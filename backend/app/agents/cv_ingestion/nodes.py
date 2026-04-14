"""
PATHS Backend — CV Ingestion Agent Nodes (LangGraph).

Each function is a node in the LangGraph workflow. All nodes receive the
CVIngestionState dict and return a partial state update dict.
"""

import os
import re
import json
import uuid
import hashlib
import logging
from typing import Dict, Any, List

from langchain_ollama import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter

from pypdf import PdfReader
from app.agents.cv_ingestion.state import CVIngestionState
from app.agents.cv_ingestion.schemas import CandidateExtraction
from app.core.config import get_settings
from app.core.database import SessionLocal
from app.db.models.candidate import Candidate
from app.db.models.cv_entities import (
    CandidateDocument, Skill, CandidateSkill, CandidateExperience,
    CandidateEducation, CandidateCertification
)
from app.repositories import graph_repo, vector_repo
from app.services.embedding_service import embed_documents

logger = logging.getLogger(__name__)
settings = get_settings()

# ────────────────────────────────────────────────────
# 1. load_document
# ────────────────────────────────────────────────────

def load_document(state: CVIngestionState) -> Dict:
    """Load and extract raw text from the uploaded document."""
    file_path = state["file_path"]
    job_id = state["job_id"]
    logger.info("[job=%s] Loading document: %s", job_id, file_path)
    try:
        raw_text = ""
        if file_path.lower().endswith(".pdf"):
            reader = PdfReader(file_path)
            for page in reader.pages:
                raw_text += (page.extract_text() or "") + "\n"
        elif file_path.lower().endswith(".docx"):
            import docx
            doc = docx.Document(file_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()

        if not raw_text.strip():
            logger.error("[job=%s] No text extracted from document", job_id)
            return {"errors": state["errors"] + ["No text could be extracted from document."], "status": "failed"}

        logger.info("[job=%s] Extracted %d chars of text", job_id, len(raw_text))
        return {"raw_text": raw_text, "stage": "extracted_text"}
    except Exception as e:
        logger.exception("[job=%s] Load document error", job_id)
        return {"errors": state["errors"] + [f"Load document error: {str(e)}"], "status": "failed"}


# ────────────────────────────────────────────────────
# 2. extract_structured_candidate_data
# ────────────────────────────────────────────────────

def _deterministic_extract(raw_text: str) -> dict:
    """Extract obvious fields deterministically: emails, phones, URLs."""
    result: dict = {}

    # Email
    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', raw_text)
    if emails:
        result["email"] = emails[0]

    # Phone
    phones = re.findall(r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\./0-9]{7,15}', raw_text)
    if phones:
        result["phone"] = phones[0].strip()

    # URLs (LinkedIn, GitHub)
    urls = re.findall(r'https?://[^\s<>"\']+', raw_text)
    result["links"] = urls

    return result


def extract_structured_candidate_data(state: CVIngestionState) -> Dict:
    """Use LLM + deterministic parsing to extract structured candidate data."""
    job_id = state["job_id"]
    logger.info("[job=%s] Extracting structured candidate data", job_id)
    try:
        # Deterministic extraction first
        deterministic = _deterministic_extract(state["raw_text"])

        llm = ChatOllama(
            model=settings.ollama_llm_model,
            base_url=settings.ollama_base_url,
            temperature=0.0,
        )
        structured_llm = llm.with_structured_output(CandidateExtraction)

        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "You are an expert HR parsing system. Extract ALL candidate information "
             "from the CV text below. Return structured JSON matching the schema exactly. "
             "Be thorough — extract every skill, experience, education entry, and certification."),
            ("human", "CV text:\n{text}\n\nExtract all candidate data.")
        ])

        chain = prompt | structured_llm
        result = chain.invoke({"text": state["raw_text"]})
        structured = result.model_dump()

        # Merge deterministic results (deterministic wins for email/phone)
        if deterministic.get("email"):
            structured["email"] = deterministic["email"]
        if deterministic.get("phone"):
            structured["phone"] = deterministic["phone"]

        logger.info("[job=%s] Extraction complete: %s, %d skills, %d experiences",
                    job_id, structured.get("full_name"), len(structured.get("skills", [])),
                    len(structured.get("experiences", [])))
        return {"structured_candidate": structured, "stage": "extracted_structure"}

    except Exception as e:
        logger.exception("[job=%s] LLM extraction error", job_id)
        # Fallback: try deterministic-only extraction
        fallback = _deterministic_extract(state["raw_text"])
        if fallback.get("email"):
            logger.warning("[job=%s] Using deterministic fallback", job_id)
            return {
                "structured_candidate": {
                    "full_name": "Unknown Candidate",
                    "email": fallback.get("email"),
                    "phone": fallback.get("phone"),
                    "skills": [],
                    "experiences": [],
                    "education": [],
                    "certifications": [],
                },
                "stage": "extracted_structure",
                "errors": state["errors"] + [f"LLM extraction failed, using fallback: {str(e)}"],
            }
        return {"errors": state["errors"] + [f"Extraction error: {str(e)}"], "status": "failed"}


# ────────────────────────────────────────────────────
# 3. normalize_entities
# ────────────────────────────────────────────────────

def normalize_entities(state: CVIngestionState) -> Dict:
    """Normalize and deduplicate extracted entities."""
    job_id = state["job_id"]
    structured = state.get("structured_candidate")
    if not structured:
        return {"errors": state["errors"] + ["No structured data to normalize."], "status": "failed"}

    logger.info("[job=%s] Normalizing entities", job_id)

    # Normalize skills: lowercase, trim, dedupe
    skills = structured.get("skills", [])
    normalized_skills = []
    seen = set()
    for s in skills:
        n_name = (s.get("name") or "").strip().lower()
        if n_name and n_name not in seen:
            seen.add(n_name)
            normalized_skills.append({
                "name": n_name,
                "category": (s.get("category") or "").strip() or None
            })

    structured["skills"] = normalized_skills

    # Clean experiences — reject empty entries
    clean_exps = []
    for exp in structured.get("experiences", []):
        company = (exp.get("company_name") or "").strip()
        title = (exp.get("title") or "").strip()
        if company and title:
            clean_exps.append({
                "company_name": company,
                "title": title,
                "start_date": (exp.get("start_date") or "").strip() or None,
                "end_date": (exp.get("end_date") or "").strip() or None,
                "description": (exp.get("description") or "").strip() or None,
            })
    structured["experiences"] = clean_exps

    # Clean education
    clean_edu = []
    for edu in structured.get("education", []):
        institution = (edu.get("institution") or "").strip()
        if institution:
            clean_edu.append({
                "institution": institution,
                "degree": (edu.get("degree") or "").strip() or None,
                "field_of_study": (edu.get("field_of_study") or "").strip() or None,
                "start_date": (edu.get("start_date") or "").strip() or None,
                "end_date": (edu.get("end_date") or "").strip() or None,
            })
    structured["education"] = clean_edu

    # Clean certifications
    clean_certs = []
    for cert in structured.get("certifications", []):
        name = (cert.get("name") or "").strip()
        if name:
            clean_certs.append({
                "name": name,
                "issuer": (cert.get("issuer") or "").strip() or None,
            })
    structured["certifications"] = clean_certs

    # Clean core fields
    structured["full_name"] = (structured.get("full_name") or "Unknown").strip()
    structured["email"] = (structured.get("email") or "").strip() or None
    structured["phone"] = (structured.get("phone") or "").strip() or None
    structured["location_text"] = (structured.get("location_text") or "").strip() or None
    structured["summary"] = (structured.get("summary") or "").strip() or None

    logger.info("[job=%s] Normalized: %d skills, %d exps, %d edu, %d certs",
                job_id, len(normalized_skills), len(clean_exps), len(clean_edu), len(clean_certs))

    return {"normalized_candidate": structured, "stage": "normalized_entities"}


# ────────────────────────────────────────────────────
# 4. persist_postgres
# ────────────────────────────────────────────────────

def persist_postgres(state: CVIngestionState) -> Dict:
    """Persist all normalized candidate data to PostgreSQL in one transaction."""
    job_id = state["job_id"]
    normalized = state.get("normalized_candidate")
    if not normalized:
        return {"errors": state["errors"] + ["No normalized data to persist."], "status": "failed"}

    candidate_id = state.get("candidate_id") or str(uuid.uuid4())
    doc_id = state.get("document_id") or str(uuid.uuid4())

    logger.info("[job=%s][candidate=%s] Persisting to PostgreSQL", job_id, candidate_id)

    db = SessionLocal()
    try:
        candidate_uuid = uuid.UUID(candidate_id)
        from sqlalchemy import select

        # ── Candidate core (upsert) ──
        candidate = db.get(Candidate, candidate_uuid)
        if not candidate:
            candidate = Candidate(
                id=candidate_uuid,
                full_name=normalized.get("full_name") or "Unknown",
                email=normalized.get("email"),
                phone=normalized.get("phone"),
                location_text=normalized.get("location_text"),
                summary=normalized.get("summary"),
                years_experience=normalized.get("years_experience"),
                status="active",
            )
            db.add(candidate)
        else:
            candidate.full_name = normalized.get("full_name") or candidate.full_name
            candidate.email = normalized.get("email") or candidate.email
            candidate.phone = normalized.get("phone") or candidate.phone
            candidate.location_text = normalized.get("location_text") or candidate.location_text
            candidate.summary = normalized.get("summary") or candidate.summary
            candidate.years_experience = normalized.get("years_experience") or candidate.years_experience
        db.flush()

        # ── Candidate document ──
        raw_text = state.get("raw_text", "")
        checksum = hashlib.sha256(raw_text.encode("utf-8")).hexdigest() if raw_text else None

        # Determine mime type
        file_path = state.get("file_path", "")
        ext = os.path.splitext(file_path)[1].lower()
        mime_map = {".pdf": "application/pdf", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".txt": "text/plain"}
        mime_type = mime_map.get(ext, "application/octet-stream")

        doc = CandidateDocument(
            id=uuid.UUID(doc_id),
            candidate_id=candidate_uuid,
            original_filename=os.path.basename(file_path),
            mime_type=mime_type,
            storage_path_or_url=file_path,
            raw_text=raw_text,
            checksum=checksum,
        )
        db.add(doc)

        # ── Skills (upsert skill catalog + link) ──
        for skill_data in normalized.get("skills", []):
            skill_name = skill_data.get("name")
            if not skill_name:
                continue

            existing_skill = db.execute(
                select(Skill).where(Skill.normalized_name == skill_name)
            ).scalar_one_or_none()
            if not existing_skill:
                existing_skill = Skill(normalized_name=skill_name, category=skill_data.get("category"))
                db.add(existing_skill)
                db.flush()

            # Check if link already exists
            existing_link = db.execute(
                select(CandidateSkill).where(
                    CandidateSkill.candidate_id == candidate_uuid,
                    CandidateSkill.skill_id == existing_skill.id,
                )
            ).scalar_one_or_none()
            if not existing_link:
                cs = CandidateSkill(candidate_id=candidate_uuid, skill_id=existing_skill.id)
                db.add(cs)

        # ── Experiences ──
        for exp in normalized.get("experiences", []):
            db.add(CandidateExperience(
                candidate_id=candidate_uuid,
                company_name=exp.get("company_name", "Unknown"),
                title=exp.get("title", "Unknown Role"),
                start_date=exp.get("start_date"),
                end_date=exp.get("end_date"),
                description=exp.get("description"),
            ))

        # ── Education ──
        for edu in normalized.get("education", []):
            db.add(CandidateEducation(
                candidate_id=candidate_uuid,
                institution=edu.get("institution", "Unknown"),
                degree=edu.get("degree"),
                field_of_study=edu.get("field_of_study"),
                start_date=edu.get("start_date"),
                end_date=edu.get("end_date"),
            ))

        # ── Certifications ──
        for cert in normalized.get("certifications", []):
            db.add(CandidateCertification(
                candidate_id=candidate_uuid,
                name=cert.get("name", "Unknown"),
                issuer=cert.get("issuer"),
            ))

        db.commit()
        logger.info("[job=%s][candidate=%s] PostgreSQL persistence complete", job_id, candidate_id)
        return {"candidate_id": candidate_id, "document_id": doc_id, "stage": "persisted_postgres"}

    except Exception as e:
        db.rollback()
        logger.exception("[job=%s] Postgres persistence error", job_id)
        return {"errors": state["errors"] + [f"Postgres error: {str(e)}"], "status": "failed"}
    finally:
        db.close()


# ────────────────────────────────────────────────────
# 5. project_to_age
# ────────────────────────────────────────────────────

def project_to_age(state: CVIngestionState) -> Dict:
    """Project all candidate relationships into the Apache AGE graph."""
    job_id = state["job_id"]
    candidate_id = state.get("candidate_id")
    document_id = state.get("document_id")
    normalized = state.get("normalized_candidate")

    logger.info("[job=%s][candidate=%s] Projecting to AGE graph", job_id, candidate_id)
    db = SessionLocal()
    try:
        graph_repo.init_graph(db)

        # Candidate node
        graph_repo.project_candidate(db, candidate_id, normalized)

        # Skills
        for skill in normalized.get("skills", []):
            graph_repo.project_skill(db, candidate_id, skill["name"], skill["name"])

        # Companies / work experience
        for exp in normalized.get("experiences", []):
            company = exp.get("company_name")
            if company:
                graph_repo.project_company(db, candidate_id, company, exp.get("title", ""))

        # Education
        for edu in normalized.get("education", []):
            institution = edu.get("institution")
            if institution:
                graph_repo.project_education(db, candidate_id, institution, edu.get("degree", ""))

        # Certifications
        for cert in normalized.get("certifications", []):
            name = cert.get("name")
            if name:
                graph_repo.project_certification(db, candidate_id, name, cert.get("issuer", ""))

        # Document linkage
        if document_id:
            graph_repo.project_document(db, candidate_id, document_id)

        db.commit()
        logger.info("[job=%s] AGE projection complete", job_id)
        return {"stage": "projected_age"}
    except Exception as e:
        db.rollback()
        logger.exception("[job=%s] AGE projection error", job_id)
        # AGE failure after PG commit: mark stage failed but allow retry
        return {"errors": state["errors"] + [f"AGE error: {str(e)}"], "stage": "age_failed"}
    finally:
        db.close()


# ────────────────────────────────────────────────────
# 6. chunk_document
# ────────────────────────────────────────────────────

def chunk_document(state: CVIngestionState) -> Dict:
    """Split the raw CV text into chunks preserving section awareness."""
    job_id = state["job_id"]
    raw_text = state.get("raw_text", "")
    logger.info("[job=%s] Chunking document (%d chars)", job_id, len(raw_text))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(raw_text)
    logger.info("[job=%s] Created %d chunks", job_id, len(chunks))
    return {"chunks": chunks, "stage": "chunked"}


# ────────────────────────────────────────────────────
# 7. embed_chunks
# ────────────────────────────────────────────────────

def embed_chunks(state: CVIngestionState) -> Dict:
    """Embed all text chunks using the local Ollama embedding model."""
    job_id = state["job_id"]
    chunks = state.get("chunks", [])
    if not chunks:
        logger.warning("[job=%s] No chunks to embed", job_id)
        return {"embeddings": [], "stage": "embedded"}
    try:
        logger.info("[job=%s] Embedding %d chunks", job_id, len(chunks))
        embeddings = embed_documents(chunks)
        logger.info("[job=%s] Embedding complete, dim=%d", job_id, len(embeddings[0]) if embeddings else 0)
        return {"embeddings": embeddings, "stage": "embedded"}
    except Exception as e:
        logger.exception("[job=%s] Embedding error", job_id)
        return {"errors": state["errors"] + [f"Embedding error: {str(e)}"], "stage": "embed_failed"}


# ────────────────────────────────────────────────────
# 8. upsert_qdrant
# ────────────────────────────────────────────────────

def upsert_qdrant(state: CVIngestionState) -> Dict:
    """Upsert embedded chunks to Qdrant vector store."""
    job_id = state["job_id"]
    chunks = state.get("chunks", [])
    embeddings = state.get("embeddings", [])
    cand_id = state.get("candidate_id")
    doc_id = state.get("document_id")

    if chunks and embeddings:
        try:
            logger.info("[job=%s] Upserting %d vectors to Qdrant", job_id, len(chunks))
            vector_repo.init_collection()
            vector_repo.upsert_chunks(cand_id, doc_id, chunks, embeddings)
            logger.info("[job=%s] Qdrant upsert complete", job_id)
            return {"stage": "upserted_qdrant"}
        except Exception as e:
            logger.exception("[job=%s] Qdrant upsert error", job_id)
            return {"errors": state["errors"] + [f"Qdrant error: {str(e)}"], "stage": "qdrant_failed"}
    logger.warning("[job=%s] No chunks/embeddings to upsert", job_id)
    return {"stage": "upserted_qdrant"}


# ────────────────────────────────────────────────────
# 9. finalize_job
# ────────────────────────────────────────────────────

def finalize_job(state: CVIngestionState) -> Dict:
    """Mark the ingestion job as completed."""
    job_id = state["job_id"]
    logger.info("[job=%s] Ingestion pipeline completed successfully", job_id)
    return {"status": "completed", "stage": "done"}


# ────────────────────────────────────────────────────
# 10. handle_failure
# ────────────────────────────────────────────────────

def handle_failure(state: CVIngestionState) -> Dict:
    """Handle pipeline failure — log errors and mark failed."""
    job_id = state["job_id"]
    errors = state.get("errors", [])
    logger.error("[job=%s] Pipeline FAILED with %d errors: %s", job_id, len(errors), "; ".join(errors))
    return {"status": "failed"}
